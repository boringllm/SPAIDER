"""Reference-document handling: extract plain text from operator-attached files.

The operator can attach documents (scope/rules-of-engagement, prior reports, target docs,
credential sheets, …) to a session so they inform the engagement. This module turns each
uploaded file into plain text that the agents can read.

Supported inputs: Markdown / plain text (stdlib), PDF (``pypdf``), Word ``.docx``
(``python-docx``). PDF/Word parsing is OPTIONAL — if the library isn't installed we return a
clear, non-fatal error string instead of raising, so the rest of the app keeps working and the
operator sees exactly what to install."""
from __future__ import annotations

import re
from pathlib import Path

# Extensions the upload endpoint accepts. Anything else is rejected with a 400.
ALLOWED_EXTS = {".md", ".markdown", ".txt", ".pdf", ".docx"}

# Max bytes accepted per uploaded file (guards the host / model context).
MAX_UPLOAD_BYTES = 20 * 1024 * 1024  # 20 MB


def safe_name(name: str) -> str:
    """Reduce an uploaded filename to a safe basename (no path traversal, no separators)."""
    base = Path(name or "").name  # strip any directory components
    base = base.replace("\\", "_").replace("/", "_").strip()
    base = re.sub(r"[^A-Za-z0-9._ \-()]", "_", base)
    return base or "document"


def _extract_pdf(data: bytes) -> tuple[str, str]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", "PDF support is not installed on the Spider host. Run: pip install pypdf"
    import io

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(p.extract_text() or "") for p in reader.pages]
    except Exception as e:  # noqa: BLE001 — corrupt/encrypted PDF, etc.
        return "", f"could not read PDF: {e}"
    text = "\n\n".join(t.strip() for t in pages if t.strip())
    if not text.strip():
        return "", "no extractable text (the PDF may be scanned images — OCR is not supported)"
    return text, ""


def _extract_docx(data: bytes) -> tuple[str, str]:
    try:
        import docx  # python-docx
    except ImportError:
        return "", "Word (.docx) support is not installed on the Spider host. Run: pip install python-docx"
    import io

    try:
        d = docx.Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        return "", f"could not read .docx: {e}"
    # Emit Markdown-ish text so a document's STRUCTURE survives: Heading-styled paragraphs
    # become Markdown headings (so a Word template's section hierarchy is preserved and the
    # report agent can reproduce it exactly), list items become bullets.
    parts: list[str] = []
    for p in d.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        style = (p.style.name if p.style else "") or ""
        if style.startswith("Heading"):
            try:
                lvl = int(style.split()[-1])
            except (ValueError, IndexError):
                lvl = 2
            parts.append("#" * min(max(lvl, 1), 6) + " " + txt)
        elif style == "Title":
            parts.append("# " + txt)
        elif style.startswith("List"):
            parts.append("- " + txt)
        else:
            parts.append(txt)
    # Include simple table content as Markdown tables (header + separator + rows).
    for table in d.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        rows = [r for r in rows if any(r)]
        if not rows:
            continue
        parts.append("")
        parts.append("| " + " | ".join(rows[0]) + " |")
        parts.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        for r in rows[1:]:
            parts.append("| " + " | ".join(r) + " |")
    text = "\n".join(parts)
    if not text.strip():
        return "", "no extractable text in the document"
    return text, ""


def extract_text(data: bytes, filename: str) -> tuple[str, str]:
    """Extract plain text from ``data`` based on ``filename``'s extension.

    Returns ``(text, error)``: on success ``error`` is empty; on failure ``text`` is empty and
    ``error`` explains why (unsupported type, missing parser, or unreadable file). Never raises."""
    ext = Path(filename).suffix.lower()
    if ext in (".md", ".markdown", ".txt"):
        return data.decode("utf-8", errors="replace"), ""
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    return "", f"unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTS))}"


# --------------------------------------------------------------------------- #
# Markdown -> Word (.docx) — render the report as a structured Word document.
# --------------------------------------------------------------------------- #
# Inline emphasis: **bold**, __bold__, *italic*, _italic_, `code`.
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`)")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBER_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")


def _add_inline_runs(paragraph, text: str) -> None:
    """Add ``text`` to a docx paragraph, honouring inline **bold**/*italic*/`code` markers."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**")) or (part.startswith("__") and part.endswith("__")):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def _looks_like_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def markdown_to_docx(md_text: str, out_path: str) -> tuple[bool, str]:
    """Render a Markdown report to a structured Word (.docx) file at ``out_path``.

    Maps the Markdown the report agent produces to native Word constructs so the Word output
    mirrors the (template-driven) structure: ``#``..``######`` -> Word Heading 1..6, ``-``/``*``
    -> bullet lists, ``1.`` -> numbered lists, ``> `` -> quote, ```` ``` ```` -> monospace code,
    ``| a | b |`` -> Word tables, and inline **bold**/*italic*/`code`. Returns ``(ok, error)`` and
    never raises; if ``python-docx`` is missing it returns a clear, non-fatal error."""
    try:
        import docx
        from docx.shared import Pt
    except ImportError:
        return False, "Word (.docx) output needs python-docx on the Spider host. Run: pip install python-docx"

    try:
        doc = docx.Document()
        lines = (md_text or "").replace("\r\n", "\n").split("\n")
        i, n = 0, len(lines)
        in_code = False
        code_buf: list[str] = []
        while i < n:
            line = lines[i]
            stripped = line.strip()

            # fenced code block
            if stripped.startswith("```"):
                if in_code:
                    para = doc.add_paragraph()
                    run = para.add_run("\n".join(code_buf))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    code_buf = []
                    in_code = False
                else:
                    in_code = True
                i += 1
                continue
            if in_code:
                code_buf.append(line)
                i += 1
                continue

            # markdown table: a header row followed by a |---|---| separator
            if _looks_like_table_row(line) and i + 1 < n and set(lines[i + 1].strip()) <= set("|-: "):
                header = _split_row(line)
                i += 2  # skip header + separator
                body: list[list[str]] = []
                while i < n and _looks_like_table_row(lines[i]):
                    body.append(_split_row(lines[i]))
                    i += 1
                table = doc.add_table(rows=1, cols=len(header))
                try:
                    table.style = "Light Grid Accent 1"
                except Exception:  # noqa: BLE001 — style may not exist in all templates
                    pass
                for j, cell in enumerate(header):
                    hc = table.rows[0].cells[j].paragraphs[0]
                    _add_inline_runs(hc, cell)
                    for r in hc.runs:
                        r.bold = True
                for brow in body:
                    cells = table.add_row().cells
                    for j in range(len(header)):
                        _add_inline_runs(cells[j].paragraphs[0], brow[j] if j < len(brow) else "")
                continue

            if not stripped:
                i += 1
                continue

            m = _HEADING_RE.match(stripped)
            if m:
                doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 6))
                i += 1
                continue
            if stripped in ("---", "***", "___"):
                i += 1
                continue
            if stripped.startswith(">"):
                para = doc.add_paragraph(style="Intense Quote") if _has_style(doc, "Intense Quote") else doc.add_paragraph()
                _add_inline_runs(para, stripped.lstrip("> ").strip())
                i += 1
                continue
            m = _BULLET_RE.match(line)
            if m:
                para = doc.add_paragraph(style="List Bullet") if _has_style(doc, "List Bullet") else doc.add_paragraph()
                _add_inline_runs(para, m.group(1))
                i += 1
                continue
            m = _NUMBER_RE.match(line)
            if m:
                para = doc.add_paragraph(style="List Number") if _has_style(doc, "List Number") else doc.add_paragraph()
                _add_inline_runs(para, m.group(1))
                i += 1
                continue
            # plain paragraph
            _add_inline_runs(doc.add_paragraph(), stripped)
            i += 1

        if in_code and code_buf:  # unterminated fence
            run = doc.add_paragraph().add_run("\n".join(code_buf))
            run.font.name = "Consolas"
            run.font.size = Pt(9)

        doc.save(out_path)
        return True, ""
    except Exception as e:  # noqa: BLE001
        return False, f"failed to build .docx: {e}"


def _has_style(doc, name: str) -> bool:
    try:
        return any(s.name == name for s in doc.styles)
    except Exception:  # noqa: BLE001
        return False
